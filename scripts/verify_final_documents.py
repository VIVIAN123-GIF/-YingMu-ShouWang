"""Run structural checks and render every final-delivery PDF page to PNG."""

from __future__ import annotations

import argparse
import json
import zipfile
from pathlib import Path

import fitz
from docx import Document


def verify_docx(path: Path) -> dict:
    document = Document(path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    if not text.strip():
        raise ValueError(f"empty DOCX: {path}")
    with zipfile.ZipFile(path) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    if 'w:type="pct"' in xml:
        raise ValueError(f"percentage table width found: {path}")
    return {
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "contains_pending": "PENDING_" in text,
    }


def verify_pdf(path: Path, render_root: Path) -> dict:
    pdf = fitz.open(path)
    if pdf.page_count < 1:
        raise ValueError(f"empty PDF: {path}")
    page_reports = []
    output_dir = render_root / path.stem
    output_dir.mkdir(parents=True, exist_ok=True)
    for stale_page in output_dir.glob("page-*.png"):
        stale_page.unlink()
    for index, page in enumerate(pdf):
        text = page.get_text().strip()
        if not text:
            raise ValueError(f"textless PDF page: {path}, page {index + 1}")
        rect = page.rect
        overflow = []
        for block in page.get_text("blocks"):
            x0, y0, x1, y1 = block[:4]
            if x0 < -1 or y0 < -1 or x1 > rect.width + 1 or y1 > rect.height + 1:
                overflow.append([round(x0, 2), round(y0, 2), round(x1, 2), round(y1, 2)])
        if overflow:
            raise ValueError(f"out-of-page text block: {path}, page {index + 1}: {overflow}")
        pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
        png = output_dir / f"page-{index + 1:03d}.png"
        pixmap.save(png)
        page_reports.append({"page": index + 1, "characters": len(text), "png": png.as_posix()})
    return {"pages": pdf.page_count, "page_reports": page_reports}


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--docx-dir", type=Path, default=Path("final-delivery/output/docx"))
    parser.add_argument("--pdf-dir", type=Path, default=Path("final-delivery/output/pdf"))
    parser.add_argument("--render-dir", type=Path, default=Path("tmp/final-document-qa"))
    parser.add_argument("--report", type=Path, default=Path("tmp/final-document-qa/report.json"))
    args = parser.parse_args()
    docx_files = sorted(args.docx_dir.glob("*.docx"))
    pdf_files = sorted(args.pdf_dir.glob("*.pdf"))
    if len(docx_files) != 8 or len(pdf_files) != 8:
        raise SystemExit(f"expected 8 DOCX and 8 PDF files, found {len(docx_files)} and {len(pdf_files)}")
    report = {"status": "PASS", "documents": {}}
    for docx_path, pdf_path in zip(docx_files, pdf_files):
        if docx_path.stem != pdf_path.stem:
            raise ValueError(f"DOCX/PDF name mismatch: {docx_path.name}, {pdf_path.name}")
        report["documents"][docx_path.stem] = {
            "docx": verify_docx(docx_path),
            "pdf": verify_pdf(pdf_path, args.render_dir),
        }
    args.report.parent.mkdir(parents=True, exist_ok=True)
    args.report.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps({
        "status": "PASS",
        "documents": len(report["documents"]),
        "pages": sum(item["pdf"]["pages"] for item in report["documents"].values()),
        "report": args.report.as_posix(),
    }, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
