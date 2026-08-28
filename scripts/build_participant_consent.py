"""Build printable two-page consent forms for P01-P03."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

from scripts.build_final_documents import (
    BLUE,
    DARK_BLUE,
    EAST_ASIA_FONT,
    LATIN_FONT,
    LIGHT_BLUE,
    MUTED,
    _add_page_field,
    _set_cell_margins,
    _set_font,
    _set_table_geometry,
    _shade,
    load_profile,
)


PARTICIPANTS = ("P01", "P02", "P03")
CONTENT_WIDTH_DXA = 9800


def _set_repeat_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    properties.append(marker)


def _paragraph(document: Document, text: str = "", *, bold: bool = False, size: float = 10.2,
               after: float = 4, align=None, color: str | None = None):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.12
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    _set_font(run, size=size, bold=bold, color=color)
    return paragraph


def _heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    _set_font(run, size=11.5, bold=True, color=BLUE)


def _checkbox(document: Document, text: str, *, required: bool = False, size: float = 9.8, after: float = 3) -> None:
    marker = "□"
    suffix = "（必选）" if required else ""
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Mm(4)
    paragraph.paragraph_format.first_line_indent = Mm(-4)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.08
    run = paragraph.add_run(f"{marker} {text}{suffix}")
    _set_font(run, size=size)


def _metadata_table(document: Document, participant_id: str, profile: dict[str, str]) -> None:
    rows = [
        ("参与者匿名编号", participant_id, "授权记录编号", f"AUTH-{participant_id}-________"),
        ("日期时间", "北京时间：________________", "保存截止", "北京时间：________________"),
        ("学校与项目", f"{profile['school']} · 萤目守望", "项目负责人", f"{profile['contact_name']}  {profile['mobile']}"),
    ]
    table = document.add_table(rows=len(rows), cols=4)
    _set_table_geometry(table, [1650, 3250, 1650, 3250])
    for row_index, values in enumerate(rows):
        for column_index, value in enumerate(values):
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell, top=100, bottom=100, start=120, end=120)
            if column_index % 2 == 0:
                _shade(cell._tc, LIGHT_BLUE)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            run = paragraph.add_run(value)
            _set_font(run, size=9.1, bold=column_index % 2 == 0)


def _signature_table(document: Document) -> None:
    rows = [
        ("参与者签名", "________________", "日期", "____年__月__日"),
        ("实验负责人签名", "________________", "日期", "____年__月__日"),
        ("现场安全员签名", "________________", "日期", "____年__月__日"),
    ]
    table = document.add_table(rows=len(rows), cols=4)
    _set_table_geometry(table, [1750, 3150, 1200, 3700])
    for row_index, row in enumerate(table.rows):
        for index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell, top=75, bottom=75, start=120, end=120)
            if index in {0, 2}:
                _shade(cell._tc, LIGHT_BLUE)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(rows[row_index][index])
            _set_font(run, size=9.5, bold=index in {0, 2})


def _page_break(document: Document) -> None:
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def build_consent(participant_id: str, profile: dict[str, str], output: Path) -> None:
    document = Document()
    section = document.sections[0]
    # Named form override: A4 is used because the forms will be printed and signed in China.
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(14)
    section.right_margin = Mm(16)
    section.bottom_margin = Mm(14)
    section.left_margin = Mm(16)
    section.header_distance = Mm(7)
    section.footer_distance = Mm(7)

    normal = document.styles["Normal"]
    normal.font.name = LATIN_FONT
    normal.font.size = Pt(10.2)
    fonts = normal.element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), LATIN_FONT)
    fonts.set(qn("w:hAnsi"), LATIN_FONT)
    fonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)

    header = section.header.paragraphs[0]
    header.text = f"吉林大学 · 萤目守望受控视频实验 · {participant_id}"
    _set_font(header.runs[0], size=8.5, color=MUTED)
    _add_page_field(section.footer.paragraphs[0])

    title = _paragraph(
        document,
        "“萤目守望”项目全流程受控视频、数据与肖像授权及安全确认",
        bold=True,
        size=16,
        after=4,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        color=DARK_BLUE,
    )
    title.paragraph_format.keep_with_next = True
    _paragraph(
        document,
        "适配XH-202617挑战杯 · 健康成年人受控工程验证 · 非医疗研究",
        bold=True,
        size=10,
        after=7,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        color=BLUE,
    )
    _metadata_table(document, participant_id, profile)

    _heading(document, "一、项目与素材说明")
    _paragraph(
        document,
        "本人自愿参加“萤目守望”居家养老风险预警系统的全流程工程验证。本项目为挑战杯揭榜挂帅XH-202617赛题参赛研发项目，由P01、P02、P03共同完成部分场景拍摄、现场记录和系统验证。三名参与者为同一家庭共同居住人员，摄像头部署于家庭居住环境并由三人共用；测试过程中可能出现其他家庭成员入镜或多人同框。",
    )
    _paragraph(document, "本人明确知悉：本测试为健康成年人受控模拟工程验证，不是老年人临床试验；项目不提供医学诊断、疾病筛查或临床诊疗服务；分时段安排仅为工程任务分配，不代表画面内仅有对应编号的一人。", after=5)

    _heading(document, "二、授权范围")
    _paragraph(
        document,
        "本授权覆盖授权有效期内本人参加的全部相关活动，不限于单次拍摄或单组动作。活动包括受控视频采集、正常活动和稳定性记录，以及为参赛作品所需的标注、复核、算法推理和系统验证。所有动作均为可控模拟，严禁真实跌倒等危险动作。",
    )
    for text in [
        "同意采集本人相关的视频画面、必要环境背景音、匿名动作标签、姿态关键点、步态和躯干特征、时间戳、质量指标、居家场景信息及系统输出。",
        "同意对授权素材进行整理、脱敏、标注、人工复核、算法推理、业务规则验证、回放、统计分析、错误分析、稳定性评估和工程性能测评。",
        "同意将授权素材用于参赛作品开发调试、系统联调、前后端测试、萤石设备接入验证、异常降级验证、研究报告、测试报告、技术文档、答辩及复赛/决赛材料。",
        "同意赛事主办方和评审专家在评审及必要原始材料抽查中查阅授权范围内的材料。",
        "同意在研究报告、测试报告、技术文档、受限提交材料和演示视频中使用匿名统计、工程指标、系统结果、脱敏截图、骨架或轮廓画面。",
        "知悉原始视频、音频、签字件和可识别截图仅存放于受控私有目录，不上传公开代码仓库、公开在线入口或未经授权的网络渠道。",
    ]:
        _checkbox(document, text, required=True, size=9.5, after=2)
    _heading(document, "三、画面展示与公开边界")
    _checkbox(document, "赛事受限展示（至少勾选一项）：□ 原始肖像  □ 人脸模糊  □ 仅骨架轮廓。", required=True, size=9.5)
    _checkbox(document, "原始肖像即使获勾选，也仅限赛事主办方或评委可访问的受限提交材料；不得用于公开网页、GitHub Pages、公开视频或社交媒体。", size=9.5)
    _checkbox(document, "公开展示仅使用人脸模糊或无法识别身份的骨架/轮廓画面；未勾选的展示形式视为不同意。", size=9.5)
    _paragraph(document, "公开在线入口不托管原始视频、签字扫描件、真实平台截图或萤石凭证，仅展示脱敏模拟数据和汇总指标。", size=9.3, after=2)
    _heading(document, "四、家庭共用摄像头与多人同框说明")
    for text in [
        "本人知晓P01、P02、P03为同一家庭，共用家庭固定摄像头，全部测试按照真实家庭环境和分时段安排开展。",
        "本人知晓稳定性时段为P01 06:00-11:00、P02 11:00-16:00、P03 16:00-21:00；时段仅为任务划分，不保证画面内只有对应编号的一人。",
        "本人知晓其他家庭成员可能偶然入镜或与本人同框。本授权仅对签署人本人有效；未签署人员的可识别画面须另行授权，或在对外提交前完成脱敏。",
        "同意项目团队客观记录多人同框，但不得将家庭单摄像头测试描述为个人隔离对照实验或个人级性能证明。",
    ]:
        _checkbox(document, text, required=True, size=9.5, after=2)

    _heading(document, "五、数据处理与材料使用边界")
    for text in [
        "知晓固定角色分工：P01用于规则校准，P02用于候选规则验证，P03用于规则冻结后的最终测试；三人均可参与拍摄、安全确认和非结果性记录。P03锁定及正式推理期间不得接触正式结果、修改标签或阈值、筛选样本或选择性重跑。",
        "知悉匿名编号仅用于工程台账，不代表个人能力、健康水平或风险等级。",
        "同意在固定版本和冻结规则下处理素材，如实记录无效、低质量、多人同框、缺失和中止原因，不通过筛选或篡改制造正向结果。",
        "知悉项目输出仅用于工程可行性和系统行为描述，不构成临床有效性、跌倒概率、疾病诊断或老年人群泛化结论。",
        "同意必要的受控备份和灾备副本，备份遵守同等访问权限、保存期限和到期删除规则。",
        "知悉原始视频、音频和可识别素材不得上传给大模型、公开在线服务或无关第三方；新增数据类型、处理用途或公开渠道须重新说明并取得书面授权。",
    ]:
        _checkbox(document, text, required=True, size=9.4, after=2)

    _heading(document, "六、保存、撤回与用途限制")
    _paragraph(document, f"允许用途：挑战杯XH-202617参赛作品开发、拍摄、标注、算法和系统测试、工程验证、研究报告、测试报告、技术文档、演示材料、受限赛事评审和必要的内部复核。保存期限最迟至{profile['retention_until']}，到期安全删除或重新取得书面授权。", size=9.4)
    _paragraph(document, "禁止用途：身份识别、生物特征识别、商业广告营销、医学诊断、保险或就业决策、对本人健康作个体化疾病结论、未经再次书面授权的公开传播或向无关第三方提供原始音视频。", size=9.4)
    _paragraph(document, "撤回后，尚未提交的报告、视频、截图、可识别中间数据及备份中的对应素材应停止使用并移除。无法反向识别个人的匿名聚合统计、完整性哈希和必要审计记录，可在不扩展用途、不重新识别个人的前提下保留并记录理由；已进入赛事评审流程或依法必须留存的副本，团队停止新增使用但无法承诺立即删除全部副本。", size=9.4)

    _heading(document, "七、安全边界与停止条件")
    for text in [
        "本人已满18周岁，测试当日无头晕、急性伤病或不适合活动的情况。",
        "现场配备防滑座椅、无障碍路线和画外安全员，可随时提供保护。",
        "严禁真实跌倒、刻意绊倒、闭眼行走、失控急转身、移除支撑物或高风险重复起身。",
        "出现头晕、疼痛、疲劳、失衡风险、碰撞风险或本人要求停止时，立即终止采集；片段标记为ABORTED，不再作为有效样本。",
        "参与者可以拒绝任一动作，正式提交前可以撤回授权，不承担任何不利后果。",
    ]:
        _checkbox(document, text, required=True, size=9.6, after=3)

    _heading(document, "八、确认与签字")
    _paragraph(document, "本人已完整阅读并理解全部授权及安全条款，获得提问和思考时间，签署完全自愿；本人确认匿名编号仅指代本人，本签字不代为其他家庭成员授权。", size=9.6)
    _paragraph(document, "项目负责人：____________________    数据管理联系人：____________________", size=9.5, after=2)
    _paragraph(document, "联系电话：________________________    隐私问题或撤回申请联系邮箱/地址：____________________________", size=9.5, after=2)
    _signature_table(document)
    _paragraph(document, "撤回申请时间：____________________    撤回方式：____________________", size=9.5, after=2)
    _paragraph(document, "处理结果：□ 原始素材已删除    □ 素材已进入赛事评审，仅停止新增使用。参与者应保留本授权书照片或副本。", size=9.5, after=0)

    document.core_properties.title = f"萤目守望参与者授权书-{participant_id}"
    document.core_properties.author = "吉林大学萤目守望项目组"
    document.core_properties.subject = "受控视频实验知情同意、肖像与数据授权"
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--profile", type=Path, default=Path("final-delivery/private-input/submission-profile.json")
    )
    parser.add_argument("--output-dir", type=Path, default=Path("output/consent-forms/docx"))
    args = parser.parse_args()
    profile = load_profile(args.profile)
    for participant_id in PARTICIPANTS:
        output = args.output_dir / f"参与者授权书_{participant_id}.docx"
        build_consent(participant_id, profile, output)
        print(f"BUILT {output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
