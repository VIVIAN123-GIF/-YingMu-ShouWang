"""Build formal DOCX drafts from the final-delivery Markdown sources."""

from __future__ import annotations

import argparse
import re
from dataclasses import dataclass
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


INK = "17324D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
PENDING_FILL = "FFF4CE"
PENDING_INK = "7A5A00"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
EAST_ASIA_FONT = "Microsoft YaHei"
LATIN_FONT = "Calibri"
CODE_FONT = "Consolas"
INLINE_PATTERN = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")


@dataclass(frozen=True)
class Preset:
    name: str
    body_alignment: WD_ALIGN_PARAGRAPH
    body_after: float
    body_line: float
    heading_spacing: tuple[tuple[float, float], tuple[float, float], tuple[float, float]]
    list_left: int
    list_hanging: int
    list_after: float
    list_line: float
    table_fill: str


PRESETS = {
    "narrative_proposal": Preset(
        "narrative_proposal", WD_ALIGN_PARAGRAPH.JUSTIFY, 8, 1.333,
        ((18, 10), (12, 6), (8, 4)), 540, 280, 4, 1.208, CALLOUT,
    ),
    "compact_reference_guide": Preset(
        "compact_reference_guide", WD_ALIGN_PARAGRAPH.LEFT, 6, 1.25,
        ((18, 10), (14, 7), (10, 5)), 540, 270, 4, 1.25, LIGHT_BLUE,
    ),
    "standard_business_brief": Preset(
        "standard_business_brief", WD_ALIGN_PARAGRAPH.LEFT, 6, 1.10,
        ((16, 8), (12, 6), (8, 4)), 720, 360, 8, 1.167, LIGHT_GRAY,
    ),
}


DOCUMENT_PRESETS = {
    "01": "narrative_proposal",
    "03": "standard_business_brief",
    "06": "standard_business_brief",
    "08": "standard_business_brief",
}


def _set_font(run, *, latin: str = LATIN_FONT, east_asia: str = EAST_ASIA_FONT, size=None,
              color: str | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = latin
    r_pr = run._element.get_or_add_rPr()
    fonts = r_pr.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)
    fonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _set_style_font(style, *, size: float, color: str = "000000", bold: bool = False) -> None:
    style.font.name = LATIN_FONT
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    r_pr = style.element.get_or_add_rPr()
    fonts = r_pr.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), LATIN_FONT)
    fonts.set(qn("w:hAnsi"), LATIN_FONT)
    fonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)


def _set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _shade(element, fill: str) -> None:
    properties = element.get_or_add_tcPr() if hasattr(element, "get_or_add_tcPr") else element.get_or_add_pPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _paragraph_left_border(paragraph, color: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color)
    borders.append(left)


def _add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    _set_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, value, end))
    tail = paragraph.add_run(" 页")
    _set_font(tail, size=9, color=MUTED)


def _add_numbering(document: Document, marker: str, *, decimal: bool, left: int, hanging: int) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if decimal else "bullet")
    level.append(num_fmt)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1." if decimal else marker)
    level.append(level_text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), str(left))
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), str(left))
    indent.set(qn("w:hanging"), str(hanging))
    p_pr.append(indent)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def _apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend((ilvl, num))
    p_pr.append(num_pr)


def _configure_document(document: Document, preset: Preset, title: str) -> dict[str, int]:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    normal = document.styles["Normal"]
    _set_style_font(normal, size=11)
    normal.paragraph_format.alignment = preset.body_alignment
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(preset.body_after)
    normal.paragraph_format.line_spacing = preset.body_line
    for index, (size, color) in enumerate(((16, BLUE), (13, BLUE), (12, DARK_BLUE)), 1):
        style = document.styles[f"Heading {index}"]
        _set_style_font(style, size=size, color=color, bold=True)
        before, after = preset.heading_spacing[index - 1]
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True

    code_style = document.styles.add_style("Yingmu Code Block", WD_STYLE_TYPE.PARAGRAPH)
    _set_style_font(code_style, size=9.2, color=INK)
    code_style.paragraph_format.left_indent = Inches(0.18)
    code_style.paragraph_format.right_indent = Inches(0.18)
    code_style.paragraph_format.space_before = Pt(4)
    code_style.paragraph_format.space_after = Pt(8)
    code_style.paragraph_format.line_spacing = 1.08
    callout_style = document.styles.add_style("Yingmu Callout", WD_STYLE_TYPE.PARAGRAPH)
    _set_style_font(callout_style, size=10.2, color=INK)
    callout_style.paragraph_format.left_indent = Inches(0.12)
    callout_style.paragraph_format.right_indent = Inches(0.12)
    callout_style.paragraph_format.space_before = Pt(6)
    callout_style.paragraph_format.space_after = Pt(8)
    callout_style.paragraph_format.line_spacing = 1.15

    header = section.header.paragraphs[0]
    running_title = title.replace("“萤目守望”", "").strip()
    header.text = f"萤目守望  |  {running_title}"
    _set_font(header.runs[0], size=8.5, color=MUTED)
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    _add_page_field(section.footer.paragraphs[0])
    document.core_properties.title = title
    document.core_properties.author = "萤目守望项目组"
    document.core_properties.subject = "2026年最终交付文档工作底稿"
    document.core_properties.keywords = "萤目守望, 居家安全, 多模态AI"
    return {
        "bullet": _add_numbering(document, "•", decimal=False, left=preset.list_left, hanging=preset.list_hanging),
        "decimal": _add_numbering(document, "", decimal=True, left=preset.list_left, hanging=preset.list_hanging),
        "check": _add_numbering(document, "☐", decimal=False, left=preset.list_left, hanging=preset.list_hanging),
    }


def _add_inline(paragraph, text: str, *, default_bold: bool = False, font_size: float | None = 11) -> None:
    position = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position:match.start()])
            _set_font(run, size=font_size, bold=default_bold)
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            _set_font(run, latin=CODE_FONT, east_asia=EAST_ASIA_FONT, size=9.5, color=DARK_BLUE)
        else:
            run = paragraph.add_run(token[2:-2])
            _set_font(run, size=font_size, bold=True)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        _set_font(run, size=font_size, bold=default_bold)


def _add_title_block(document: Document, title: str, lead_lines: list[str], *, cover: bool) -> None:
    if cover:
        spacer = document.add_paragraph()
        spacer.paragraph_format.space_after = Pt(100)
        kicker = document.add_paragraph()
        kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = kicker.add_run("专项研究报告")
        _set_font(run, size=11, color=BLUE, bold=True)
        kicker.paragraph_format.space_after = Pt(18)
        title_p = document.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.paragraph_format.space_after = Pt(14)
        if "前置预警" in title:
            first, second = title.split("前置预警", 1)
            run = title_p.add_run(first.rstrip())
            _set_font(run, size=24, color=INK, bold=True)
            run.add_break()
            run = title_p.add_run("前置预警" + second)
            _set_font(run, size=24, color=INK, bold=True)
        else:
            run = title_p.add_run(title)
            _set_font(run, size=25, color=INK, bold=True)
        for index, line in enumerate(lead_lines):
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(8 if index == 0 else 4)
            run = paragraph.add_run(line)
            _set_font(run, size=12 if index == 0 else 9.5, color=DARK_BLUE if index == 0 else MUTED,
                      italic=index > 0)
        date_p = document.add_paragraph()
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_p.paragraph_format.space_before = Pt(72)
        run = date_p.add_run(f"工作底稿  |  {date.today().isoformat()}")
        _set_font(run, size=10, color=MUTED)
        date_p.add_run().add_break(WD_BREAK.PAGE)
        return

    title_p = document.add_paragraph()
    title_p.paragraph_format.space_after = Pt(8)
    if "与开源软件清单" in title:
        first, second = title.split("与开源软件清单", 1)
        run = title_p.add_run(first.rstrip())
        _set_font(run, size=22, color=INK, bold=True)
        run.add_break()
        run = title_p.add_run("与开源软件清单" + second)
        _set_font(run, size=22, color=INK, bold=True)
    else:
        run = title_p.add_run(title)
        _set_font(run, size=23, color=INK, bold=True)
    for line in lead_lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(8)
        _add_inline(paragraph, line)
        _shade(paragraph._p, CALLOUT)
        _paragraph_left_border(paragraph, BLUE)


def _parse_table(lines: list[str]) -> list[list[str]]:
    rows = [[cell.strip() for cell in line.strip().strip("|").split("|")] for line in lines]
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell) for cell in rows[1]):
        del rows[1]
    return rows


def _set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_pr = table._tbl.tblPr
    width = table_pr.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
        table_pr.append(width)
    width.set(qn("w:w"), str(sum(widths)))
    width.set(qn("w:type"), "dxa")
    indent = table_pr.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        table_pr.append(indent)
    indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
    indent.set(qn("w:type"), "dxa")
    layout = table_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        table_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(value))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")


def _table_widths(rows: list[list[str]]) -> list[int]:
    columns = len(rows[0])
    maxima = [max(len(row[index]) if index < len(row) else 0 for row in rows) for index in range(columns)]
    weights = [max(5, min(length, 45)) for length in maxima]
    total = sum(weights)
    widths = [round(TABLE_WIDTH_DXA * weight / total) for weight in weights]
    widths[-1] += TABLE_WIDTH_DXA - sum(widths)
    return widths


def _add_table(document: Document, rows: list[list[str]], preset: Preset) -> None:
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    widths = _table_widths(rows)
    _set_table_geometry(table, widths)
    for row_index, row in enumerate(rows):
        for column_index, value in enumerate(row):
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell)
            if row_index == 0:
                _shade(cell._tc, preset.table_fill)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(value) <= 16 else WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.08
            _add_inline(paragraph, value, default_bold=row_index == 0)
            for run in paragraph.runs:
                if run.font.size is None or run.font.size.pt > 9.5:
                    run.font.size = Pt(9.5)
    after = document.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def _render_markdown(document: Document, lines: list[str], preset: Preset, numbering: dict[str, int]) -> None:
    index = 0
    active_decimal_id: int | None = None
    compact_references = False
    while index < len(lines):
        raw = lines[index]
        line = raw.strip()
        if not line:
            index += 1
            continue
        if line.startswith("```"):
            active_decimal_id = None
            code_lines = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index].rstrip())
                index += 1
            paragraph = document.add_paragraph(style="Yingmu Code Block")
            _shade(paragraph._p, LIGHT_GRAY)
            run = paragraph.add_run("\n".join(code_lines))
            _set_font(run, latin=CODE_FONT, east_asia=EAST_ASIA_FONT, size=9.2, color=INK)
            index += 1
            continue
        if line.startswith("|"):
            active_decimal_id = None
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            _add_table(document, _parse_table(table_lines), preset)
            continue
        heading = re.match(r"^(#{2,4})\s+(.+)$", line)
        if heading:
            active_decimal_id = None
            level = min(3, len(heading.group(1)) - 1)
            paragraph = document.add_paragraph(style=f"Heading {level}")
            _add_inline(paragraph, heading.group(2), default_bold=True, font_size=None)
            compact_references = heading.group(2) == "参考资料"
            index += 1
            continue
        if line.startswith("> "):
            active_decimal_id = None
            paragraph = document.add_paragraph(style="Yingmu Callout")
            content = line[2:]
            _add_inline(paragraph, content, default_bold="PENDING_" in content)
            _shade(paragraph._p, PENDING_FILL if "PENDING_" in content else CALLOUT)
            _paragraph_left_border(paragraph, PENDING_INK if "PENDING_" in content else BLUE)
            index += 1
            continue
        numbered = re.match(r"^\d+\.\s+(.+)$", line)
        bullet = re.match(r"^-\s+(.+)$", line)
        if numbered or bullet:
            content = (numbered or bullet).group(1)
            check = content.startswith("[ ] ")
            if check:
                content = content[4:]
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(2 if compact_references else preset.list_after)
            paragraph.paragraph_format.line_spacing = 1.0 if compact_references else preset.list_line
            if numbered:
                if active_decimal_id is None:
                    active_decimal_id = _add_numbering(
                        document, "", decimal=True, left=preset.list_left, hanging=preset.list_hanging
                    )
                num_id = active_decimal_id
            else:
                active_decimal_id = None
                num_id = numbering["check" if check else "bullet"]
            _apply_numbering(paragraph, num_id)
            _add_inline(paragraph, content, font_size=9 if compact_references else 11)
            index += 1
            continue
        active_decimal_id = None
        paragraph = document.add_paragraph()
        _add_inline(paragraph, line)
        index += 1


def build_document(source: Path, output: Path) -> None:
    lines = source.read_text(encoding="utf-8").splitlines()
    if not lines or not lines[0].startswith("# "):
        raise ValueError(f"missing document title: {source}")
    title = lines[0][2:].strip()
    prefix = source.name[:2]
    preset = PRESETS[DOCUMENT_PRESETS.get(prefix, "compact_reference_guide")]
    document = Document()
    numbering = _configure_document(document, preset, title)
    lead_lines = []
    body_start = 1
    while body_start < len(lines) and (not lines[body_start].strip() or not lines[body_start].startswith("## ")):
        if lines[body_start].strip():
            lead_lines.append(lines[body_start].strip())
        body_start += 1
    _add_title_block(document, title, lead_lines, cover=prefix == "01")
    _render_markdown(document, lines[body_start:], preset, numbering)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--source-dir", type=Path, default=Path("final-delivery/docs"))
    parser.add_argument("--output-dir", type=Path, default=Path("final-delivery/output/docx"))
    args = parser.parse_args()
    sources = sorted(args.source_dir.glob("*.md"))
    if len(sources) != 8:
        raise SystemExit(f"expected 8 Markdown sources, found {len(sources)}")
    for source in sources:
        output = args.output_dir / f"{source.stem}.docx"
        build_document(source, output)
        print(f"BUILT {output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
